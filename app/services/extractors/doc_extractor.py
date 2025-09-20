# app/services/extractors/doc_extractor.py
from io import BytesIO
from docx import Document
from .base_extractor import BaseExtractor

class DocExtractor(BaseExtractor):
    def extract_text(self, file_bytes: BytesIO):
        """
        Extracts all text and tables from a DOCX file.
        Args:
            file_bytes (BytesIO): The DOCX file in memory.
        Returns:
            tuple:
                - full_text: str (all paragraph text)
                - tables: list of dicts with 'headers' and 'rows'
        """
        doc_stream = BytesIO(file_bytes.read())
        doc = Document(doc_stream)

        # Extract full text
        full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip() != ""])

        # Extract tables
        tables = []
        for table in doc.tables:
            # Use first row as headers
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            rows = []
            for row in table.rows[1:]:
                row_data = [cell.text.strip() for cell in row.cells]
                rows.append(row_data)
            tables.append({"headers": headers, "rows": rows})

        return full_text, "Word"
